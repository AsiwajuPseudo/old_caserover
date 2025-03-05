import csv
import json
import io
import os
import fitz
import docx
import time
import random
import openpyxl
from langchain.text_splitter import TokenTextSplitter

class Collector:
    @staticmethod
    def collect_csv(filename):
        try:
            data = []
            with open(filename, 'r', newline='') as csvfile:
                reader = csv.DictReader(csvfile)
                for row in reader:
                    data.append({key: row[key] for key in row})
            return data
        except FileNotFoundError:
            print(f"Error: File '{filename}' not found.")
            return []
        except Exception as e:
            print(f"An error occurred: {e}")
            return []

    @staticmethod
    def find_arrays(d):
        array=[]
        for key, value in d.items():
            if isinstance(value, list):
                array=value
                break
            elif isinstance(value, dict):
                array = find_arrays(value)
        return array

    @staticmethod
    def collect_json(filename):
        try:
            with open(filename, 'r') as jsonfile:
                data = json.load(jsonfile)
                updated_data=Collector.find_arrays(data)
            return updated_data
        except FileNotFoundError:
            print(f"Error: File '{filename}' not found.")
            return []
        except json.JSONDecodeError:
            print(f"Error: File '{filename}' is not a valid JSON.")
            return []
        except Exception as e:
            print(f"An error occurred: {e}")
            return []

    @staticmethod
    def collect_xlsx(filename):
        try:
            data = []
            workbook = openpyxl.load_workbook(filename)
            sheet = workbook.active
            
            headers = [cell.value for cell in sheet[1]]
            for row in sheet.iter_rows(min_row=2, values_only=True):
                data.append({headers[i]: row[i] for i in range(len(headers))})
            
            return data
        except FileNotFoundError:
            print(f"Error: File '{filename}' not found.")
            return []
        except Exception as e:
            print(f"An error occurred: {e}")
            return []

    @staticmethod
    def pdf_lines(path):
        try:
            #open the file
            pdf_document = fitz.open(path)
            document = []
            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                text = page.get_text()
                lines = text.split('\n')
                page_lines = [{'n': i + 1, 'text': line} for i, line in enumerate(lines)]
                document.append({'page': page_num + 1, 'lines': page_lines})

            return document
        except FileNotFoundError:
            print(f"Error: File '{path}' not found.")
            return []
        except Exception as e:
            print(f"An error occurred: {e}")
            return []

    @staticmethod
    def collect_docx(path):
        try:
            #open the file
            docx_file = docx.Document(path)
            paragraphs = [{'text':paragraph.text} for paragraph in docx_file.paragraphs]

            return paragraphs
        except FileNotFoundError:
            print(f"Error: File '{path}' not found.")
            return []
        except Exception as e:
            print(f"An error occurred: {e}")
            return []

    @staticmethod
    def docx_lines(path):
        try:
            #open the file
            docx_file = docx.Document(path)
            paragraphs = [paragraph.text for paragraph in docx_file.paragraphs]
            n=1
            i=0
            num=1
            doc=[]
            page=[]
            for p in paragraphs:
                if n<12 and i<len(paragraphs)-1:
                    page.append({'n':n,'text':p})
                    n=n+1
                elif i==len(paragraphs)-1:
                    page.append({'n':n,'text':p})
                    doc.append({'page':num,'lines':page})
                else:
                    doc.append({'page':num,'lines':page})
                    page=[]
                    page.append({'n':n,'text':p})
                    n=1
                    num=num+1
                i=i+1

            return doc
        except FileNotFoundError:
            print(f"Error: File '{path}' not found.")
            return []
        except Exception as e:
            print(f"An error occurred: {e}")
            return []

    @staticmethod
    def docx_styles(path):
        # Open the document
        doc = docx.Document(path)
        
        # List to store paragraphs with their styles and text
        paragraphs_with_styles = []
        
        # Iterate through each paragraph in the document
        for para in doc.paragraphs:
            paragraph_data = {
                'style': para.style.name,
                'text': para.text
            }
            paragraphs_with_styles.append(paragraph_data)
        
        return paragraphs_with_styles

    @staticmethod
    def pdf_raw(path):
        try:
            #open the file
            pdf_document = fitz.open(path)
            dataset = []
            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                text= page.get_text()
                dataset.append({'page_number':page_num,'text':text})

            return dataset
        except FileNotFoundError:
            print(f"Error: File '{path}' not found.")
            return []
        except Exception as e:
            print(f"An error occurred: {e}")
            return []

# Example usage:
# collector = Collector()
# collected_data = collector.collect_csv('data.csv')
# print(collected_data)